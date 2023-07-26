from docx import Document
from googletrans import Translator

def translate_text(text, source_language='en', target_language='ar'):
    try:
        # Initialize the translator
        translator = Translator()

        # Translate the text to the target language
        translated_text = translator.translate(text, src=source_language, dest=target_language)

        if translated_text.text:
            return translated_text.text
        else:
            print("Translation failed for:", text)
            return text

    except Exception as e:
        print("An error occurred during translation:", str(e))
        return text

def translate_docx(input_file, output_file, source_language='en', target_language='ar'):
    try:
        # Load the input Word document
        doc = Document(input_file)

        # Translate the content of each paragraph in the document
        for paragraph in doc.paragraphs:
            translated_text = translate_text(paragraph.text, source_language, target_language)
            paragraph.text = translated_text

        # Translate the content of each table in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    translated_text = translate_text(cell.text, source_language, target_language)
                    cell.text = translated_text

        # Save the translated content to the output file
        doc.save(output_file)

        print("Translation completed. The translated content is saved in:", output_file)

    except Exception as e:
        print("An error occurred:", str(e))

if __name__ == "__main__":
    # Replace 'input_file.docx' with the path to your input Word file.
    # The translated content will be saved to 'output_file_arabic.docx'.
    translate_docx('input_file.docx', 'output_file_arabic.docx')
